#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
██╗░░░██╗░█████╗░██████╗░░██████╗███████╗
██║░░░██║██╔══██╗██╔══██╗██╔════╝██╔════╝
╚██╗░██╔╝███████║██████╦╝╚█████╗░█████╗░░
░╚████╔╝░██╔══██║██╔══██╗░╚═══██╗██╔══╝░░
░░╚██╔╝░░██║░░██║██████╦╝██████╔╝███████╗
░░░╚═╝░░░╚═╝░░╚═╝╚═════╝░╚═════╝░╚══════╝

Criado por: Borge Levisberg
Versão PRO - Conversor PDF ↔ Word automático
Feito com paixão no Termux, pra hackers com fé e café.
"""
import os
import sys
import subprocess
import argparse
from pathlib import Path

# Lista de dependências necessárias
REQUIRED_PACKAGES = {
    "PyPDF2": "pypdf",
    "docx": "python-docx",
    "docx2pdf": "docx2pdf"
}

# Verifica e instala dependências
def ensure_dependencies():
    for module_name, package_name in REQUIRED_PACKAGES.items():
        try:
            __import__(module_name)
        except ImportError:
            print(f"[!] {module_name} não encontrado. Instalando {package_name}...")
            subprocess.check_call([sys.executable, "-m", "pip", "install", package_name])

# Chama antes de tudo
ensure_dependencies()

# Importações seguras agora
from PyPDF2 import PdfReader
from docx import Document
from docx2pdf import convert as docx2pdf_convert


def pdf_to_word(pdf_path, output_path):
    print(f"[+] Convertendo PDF para DOCX: {pdf_path}")
    reader = PdfReader(pdf_path)
    doc = Document()

    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)
        else:
            doc.add_paragraph(f"[Página {i+1} sem texto extraível]")

    doc.save(output_path)
    print(f"[✓] Salvo como: {output_path}")


def word_to_pdf(docx_path, output_path):
    print(f"[+] Convertendo DOCX para PDF: {docx_path}")
    try:
        docx2pdf_convert(docx_path, output_path)
        print(f"[✓] PDF gerado: {output_path}")
    except Exception as e:
        print("[!] Falha na conversão. Verifique se você tem o Microsoft Word (Windows) ou LibreOffice (Linux).")
        print("Erro:", e)


def main():
    parser = argparse.ArgumentParser(description="Conversor de PDF ↔ DOCX Profissional")
    parser.add_argument("--input", "-i", required=True, help="Arquivo de entrada (.pdf ou .docx)")
    parser.add_argument("--output", "-o", help="Arquivo de saída (opcional)")
    parser.add_argument("--to-pdf", action="store_true", help="Converte DOCX para PDF")
    parser.add_argument("--to-docx", action="store_true", help="Converte PDF para DOCX")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print("[x] Arquivo de entrada não encontrado.")
        return

    ext = input_path.suffix.lower()
    output_path = Path(args.output) if args.output else None

    if args.to_docx:
        if ext != ".pdf":
            print("[x] Para usar --to-docx, o arquivo precisa ser PDF.")
            return
        out = output_path or input_path.with_suffix(".docx")
        pdf_to_word(str(input_path), str(out))

    elif args.to_pdf:
        if ext != ".docx":
            print("[x] Para usar --to-pdf, o arquivo precisa ser DOCX.")
            return
        out = output_path or input_path.with_suffix(".pdf")
        word_to_pdf(str(input_path), str(out))

    else:
        print("[x] Você precisa especificar uma ação: --to-pdf ou --to-docx")


if __name__ == "__main__":
    main()